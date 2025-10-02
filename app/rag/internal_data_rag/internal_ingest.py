# app/rag/internal_data_rag/internal_ingest.py
# -*- coding: utf-8 -*-
# 문서 임베딩 및 ChromaDB 저장 서비스
# (MD 정규화 1단계 + XLSX "시트 단위" 청킹/메타 보존 + 메타 확장)
import os
import sys
import time
import zipfile
import re
import hashlib
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any
from collections import deque

import pdfplumber
import docx
import openpyxl
from dotenv import load_dotenv

import chromadb
from chromadb.config import Settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI

load_dotenv()

# -------------------- 환경 변수 --------------------
CHROMA_PATH = os.getenv("CHROMA_PATH", "./chroma_data")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "inside_data1")

# 기본(고정) 청킹 파라미터 — ②단계에서 적응형으로 교체 예정
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "150"))

# XLSX 보호 옵션
XLSX_MAX_ROWS_PER_SHEET = int(os.getenv("XLSX_MAX_ROWS_PER_SHEET", "10000"))
XLSX_MAX_COLS_PER_SHEET = int(os.getenv("XLSX_MAX_COLS_PER_SHEET", "512"))
XLSX_SKIP_HIDDEN_SHEETS = os.getenv("XLSX_SKIP_HIDDEN_SHEETS", "true").lower() == "true"

# 임베딩 배치 제한
EMBED_MAX_TOKENS_PER_REQUEST = int(os.getenv("EMBED_MAX_TOKENS_PER_REQUEST", "280000"))
EMBED_MAX_ITEMS_PER_REQUEST = int(os.getenv("EMBED_MAX_ITEMS_PER_REQUEST", "256"))

# 🔧 MD 정규화 옵션
NORMALIZE_TO_MD = os.getenv("NORMALIZE_TO_MD", "true").lower() == "true"
PDF_ADD_PAGE_MARKERS = os.getenv("PDF_ADD_PAGE_MARKERS", "true").lower() == "true"
PDF_HEADING_HEURISTIC = os.getenv("PDF_HEADING_HEURISTIC", "true").lower() == "true"
XLSX_SAMPLE_TOP = int(os.getenv("XLSX_SAMPLE_TOP", "50"))
XLSX_SAMPLE_BOTTOM = int(os.getenv("XLSX_SAMPLE_BOTTOM", "50"))
XLSX_SCHEMA_SCAN_ROWS = int(os.getenv("XLSX_SCHEMA_SCAN_ROWS", "200"))

# tiktoken(선택)
try:
    import tiktoken
    _TIKTOKEN_ENC = tiktoken.get_encoding("cl100k_base")
except Exception:
    _TIKTOKEN_ENC = None

client = OpenAI()

# Chroma Cloud 경고(선택)
if os.getenv("CHROMA_API_KEY") is None:
    print("경고: CHROMA_API_KEY 환경 변수가 설정되지 않았습니다.")
if os.getenv("CHROMA_TENANT") is None:
    print("경고: CHROMA_TENANT 환경 변수가 설정되지 않았습니다.")
if os.getenv("CHROMA_DATABASE") is None:
    print("경고: CHROMA_DATABASE 환경 변수가 설정되지 않았습니다.")

# 시트 블록 식별용
SHEET_HEADER_RE = re.compile(r'^### \[Sheet\] (?P<name>.+)$', re.MULTILINE)
PAGE_MARK_RE = re.compile(r'^\[PAGE\s+(\d+)\]$', re.MULTILINE)

# -------------------- 유틸 --------------------
def _detect_office_kind(path: Path) -> Optional[str]:
    try:
        if not zipfile.is_zipfile(path):
            return None
        with zipfile.ZipFile(path) as z:
            names = set(z.namelist())
        if any(n.startswith("word/") for n in names):
            return "docx"
        if any(n.startswith("xl/") for n in names):
            return "xlsx"
        return None
    except Exception:
        return None

def _estimate_tokens(text: str) -> int:
    if _TIKTOKEN_ENC is not None:
        try:
            return len(_TIKTOKEN_ENC.encode(text))
        except Exception:
            pass
    return max(1, len(text) // 4)

def _sha256(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()

def embed_texts_batched(texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
    batches: List[List[str]] = []
    current: List[str] = []
    current_tokens = 0
    for t in texts:
        tk = _estimate_tokens(t)
        if tk > EMBED_MAX_TOKENS_PER_REQUEST:
            if current:
                batches.append(current); current=[]; current_tokens=0
            batches.append([t]); continue
        if current and (current_tokens + tk > EMBED_MAX_TOKENS_PER_REQUEST or len(current) >= EMBED_MAX_ITEMS_PER_REQUEST):
            batches.append(current); current=[]; current_tokens=0
        current.append(t); current_tokens += tk
    if current:
        batches.append(current)

    all_embeddings: List[List[float]] = []
    for i, batch in enumerate(batches, 1):
        print(f"  🔎 임베딩 배치 {i}/{len(batches)} (items={len(batch)}) 요청 중...")
        resp = client.embeddings.create(model="text-embedding-3-small", input=batch)
        all_embeddings.extend([d.embedding for d in resp.data])
    return all_embeddings

# -------------------- MD 정규화 유틸 --------------------
def _escape_md_cell(text: Any) -> str:
    s = "" if text is None else str(text)
    return s.replace("|", "\\|").replace("\n", " ").strip()

def _docx_heading_level(paragraph: docx.text.paragraph.Paragraph) -> Optional[int]:
    try:
        name = (paragraph.style and paragraph.style.name) or ""
        m = re.search(r"Heading\s*([1-6])", str(name), re.IGNORECASE)
        if m:
            return int(m.group(1))
    except Exception:
        pass
    return None

def _is_pdf_heading(line: str) -> bool:
    if not line or len(line) < 6 or len(line) > 120:
        return False
    if re.match(r"^\d+(?:\.\d+)*\s+\S+", line):
        return True
    letters = [ch for ch in line if ch.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for ch in letters if ch.isupper()) / len(letters)
    return upper_ratio >= 0.6

def _infer_col_stats(values: List[Any]) -> Tuple[str, float, float]:
    non_null = [v for v in values if v not in (None, "")]
    null_rate = 1.0 - (len(non_null) / max(1, len(values)))
    def as_type(v: Any) -> str:
        if v is None or v == "":
            return "empty"
        if hasattr(v, "isoformat"):
            return "date"
        if isinstance(v, bool):
            return "bool"
        if isinstance(v, int):
            return "int"
        if isinstance(v, float):
            return "float"
        s = str(v).strip()
        if re.match(r"^\d{4}-\d{2}-\d{2}$", s):
            return "date"
        if re.match(r"^-?\d+$", s):
            return "int"
        if re.match(r"^-?\d+\.\d+$", s):
            return "float"
        return "text"
    type_counts: Dict[str, int] = {}
    for v in non_null:
        t = as_type(v)
        type_counts[t] = type_counts.get(t, 0) + 1
    inferred = max(type_counts.items(), key=lambda x: x[1])[0] if type_counts else "text"
    uniq_rate = (len(set(map(lambda x: str(x), non_null))) / max(1, len(non_null))) if non_null else 0.0
    return inferred, null_rate, uniq_rate

def _rows_to_md_table(header: List[Any], rows: List[List[Any]]) -> List[str]:
    header_e = [_escape_md_cell(h) for h in header]
    lines = ["| " + " | ".join(header_e) + " |", "| " + " | ".join(["---"] * len(header_e)) + " |"]
    for r in rows:
        cells = [_escape_md_cell(r[i] if i < len(header_e) else "") for i in range(len(header_e))]
        lines.append("| " + " | ".join(cells) + " |")
    return lines

# -------------------- Ingest 서비스 --------------------
class IngestService:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP, separators=["\n\n", "\n", " ", ""]
        )
        self.supported_extensions = {".pdf", ".docx", ".xlsx", ".csv", ".txt"}

    # ---------- 원문→Markdown ----------
    def _docx_to_md(self, path: Path) -> Tuple[str, Dict]:
        try:
            d = docx.Document(str(path))
        except Exception as e:
            raise ValueError(f"DOCX 로드 실패: {type(e).__name__}: {e}")
        lines: List[str] = []
        toc: List[str] = []
        h_count = 0
        tbl_count = 0

        for p in d.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            lvl = _docx_heading_level(p)
            if lvl:
                h_count += 1
                lines.append("#" * lvl + " " + text)
                if lvl <= 3:
                    if lvl == 1:
                        toc.append(text)
                    elif lvl == 2:
                        toc.append((toc[-1] + ">" if toc else "") + text)
                    else:
                        base = toc[-1] if toc else ""
                        toc.append((base + ">" if base else "") + text)
                continue
            lines.append(text)

        for table in d.tables:
            tbl_count += 1
            rows = []
            for row in table.rows:
                cells = [_escape_md_cell(c.text) for c in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            header = rows[0]
            sep = ["---"] * len(header)
            lines.append("")
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(sep) + " |")
            for r in rows[1:]:
                lines.append("| " + " | ".join(r) + " |")
            lines.append("")

        md_text = "\n\n".join(lines).strip()
        meta = {
            "source_type": "docx",
            "stats": {"headings": h_count, "tables": tbl_count, "sheets": 0, "pages": 0, "empty_rate": 0.0},
            "toc": toc,
            "normalizer_ver": "md-v1",
        }
        return md_text, meta

    def _pdf_to_md(self, path: Path) -> Tuple[str, Dict]:
        lines: List[str] = []
        pages = 0
        empty_pages = 0
        h_count = 0
        try:
            with pdfplumber.open(str(path)) as pdf:
                for idx, page in enumerate(pdf.pages, start=1):
                    pages += 1
                    text = page.extract_text() or ""
                    text = text.replace("\r", "\n")
                    text = re.sub(r"\n{3,}", "\n\n", text)
                    if PDF_ADD_PAGE_MARKERS:
                        lines.append("\n---\n[PAGE %d]\n---\n" % idx)
                    if not text.strip():
                        empty_pages += 1
                        continue
                    for raw_line in text.split("\n"):
                        line = raw_line.strip()
                        if not line:
                            continue
                        if PDF_HEADING_HEURISTIC and _is_pdf_heading(line):
                            h_count += 1
                            lines.append("## " + line)
                        else:
                            lines.append(line)
        except Exception as e:
            raise ValueError(f"PDF 로드 실패: {type(e).__name__}: {e}")
        md_text = "\n".join(lines).strip()
        empty_rate = (empty_pages / pages) if pages else 0.0
        meta = {
            "source_type": "pdf",
            "stats": {"headings": h_count, "tables": 0, "sheets": 0, "pages": pages, "empty_rate": round(empty_rate, 3)},
            "toc": [],
            "normalizer_ver": "md-v1",
        }
        return md_text, meta

    def _xlsx_to_md_schema_sample(self, path: Path) -> Tuple[str, Dict]:
        try:
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        except Exception as e:
            raise ValueError(f"엑셀 로드 실패: {type(e).__name__}: {e}")
        if not wb.worksheets:
            raise ValueError("엑셀에 워크시트가 없습니다.")

        lines: List[str] = []
        sheet_count = 0
        for ws in wb.worksheets:
            try:
                if XLSX_SKIP_HIDDEN_SHEETS and getattr(ws, "sheet_state", "visible") != "visible":
                    continue
            except Exception:
                pass
            sheet_count += 1
            lines.append(f"### [Sheet] {ws.title}")

            iter_kwargs = {"values_only": True}
            if XLSX_MAX_COLS_PER_SHEET and XLSX_MAX_COLS_PER_SHEET > 0:
                iter_kwargs["max_col"] = XLSX_MAX_COLS_PER_SHEET

            header: Optional[List[str]] = None
            top_rows: List[List[Any]] = []
            bottom_rows: deque = deque(maxlen=XLSX_SAMPLE_BOTTOM)
            schema_sample_rows: List[List[Any]] = []

            row_idx = 0
            for row in ws.iter_rows(**iter_kwargs):
                row_idx += 1
                last = -1
                for i, v in enumerate(row):
                    sv = (str(v).strip() if v is not None else "")
                    if sv != "":
                        last = i
                if last < 0:
                    continue
                r = list(row[: last + 1])

                if header is None:
                    header = [str(c).strip() if c is not None else f"col{i+1}" for i, c in enumerate(r)]
                    continue

                if len(top_rows) < XLSX_SAMPLE_TOP:
                    top_rows.append(r)
                bottom_rows.append(r)
                if len(schema_sample_rows) < XLSX_SCHEMA_SCAN_ROWS:
                    schema_sample_rows.append(r)
                if row_idx > XLSX_MAX_ROWS_PER_SHEET:
                    break

            if header is None:
                lines.append("(empty sheet)")
                continue

            lines.append("- columns:")
            col_count = len(header)
            for ci in range(col_count):
                col_name = header[ci] or f"col{ci+1}"
                vals = []
                for r in schema_sample_rows:
                    v = r[ci] if ci < len(r) else None
                    vals.append(v)
                inferred, null_rate, uniq_rate = _infer_col_stats(vals)
                lines.append(f"  - {col_name}: {inferred} (null {null_rate:.1%}, unique {uniq_rate:.1%})")

            if top_rows:
                lines.append("\n**sample (top %d)**" % len(top_rows))
                lines.extend(_rows_to_md_table(header, top_rows))
            if bottom_rows:
                bot_rows = list(bottom_rows)
                lines.append("\n**sample (bottom %d)**" % len(bot_rows))
                lines.extend(_rows_to_md_table(header, bot_rows))
            lines.append("")

        md_text = "\n".join(lines).strip()
        meta = {
            "source_type": "xlsx",
            "stats": {"headings": 0, "tables": 0, "sheets": sheet_count, "pages": 0, "empty_rate": 0.0},
            "toc": [],
            "normalizer_ver": "md-v1",
        }
        return md_text, meta

    def to_markdown(self, file_path: str, verbose: bool = True) -> Tuple[str, Dict]:
        p = Path(file_path)
        ext = p.suffix.lower()
        if verbose:
            print(f"  📝 Markdown 정규화: {p.name} ({ext})")
        actual = _detect_office_kind(p)
        try:
            if ext == ".pdf":
                return self._pdf_to_md(p)
            if ext == ".docx" or (actual == "docx" and ext != ".xlsx"):
                if verbose and ext != ".docx" and actual == "docx":
                    print("  ⚠️ 확장자와 다른 실제 포맷(docx) 감지 → docx→md 변환")
                return self._docx_to_md(p)
            if ext == ".xlsx" or (actual == "xlsx" and ext != ".docx"):
                if verbose and ext != ".xlsx" and actual == "xlsx":
                    print("  ⚠️ 확장자와 다른 실제 포맷(xlsx) 감지 → xlsx→md 변환")
                return self._xlsx_to_md_schema_sample(p)

            if actual == "docx":
                return self._docx_to_md(p)
            if actual == "xlsx":
                return self._xlsx_to_md_schema_sample(p)

            if verbose:
                print(f"  ⚠️ 지원하지 않는 파일 형식: {ext} (실제 포맷 미확인)")
            return "", {"source_type": "unknown", "stats": {}, "toc": [], "normalizer_ver": "md-v1"}
        except Exception as e:
            if verbose:
                print(f"  ❌ Markdown 변환 오류 ({p.name}): {e}")
            return "", {"source_type": "error", "stats": {}, "toc": [], "normalizer_ver": "md-v1"}

    # ---------- 백업 파서(미사용 경로) ----------
    def read_pdf(self, path: Path) -> str:
        texts = []
        try:
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        texts.append(t)
        except Exception as e:
            raise ValueError(f"PDF 로드 실패: {type(e).__name__}: {e}")
        return "\n\n".join(texts)

    def read_docx(self, path: Path) -> str:
        try:
            d = docx.Document(str(path))
        except Exception as e:
            raise ValueError(f"DOCX 로드 실패: {type(e).__name__}: {e}")
        acc: List[str] = []
        acc.extend([p.text for p in d.paragraphs if p.text and p.text.strip()])
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    acc.append(" | ".join(cells))
        return "\n".join(acc)

    def read_xlsx(self, path: Path) -> str:
        try:
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        except Exception as e:
            raise ValueError(f"엑셀 로드 실패: {type(e).__name__}: {e}")
        if not wb.worksheets:
            raise ValueError("엑셀에 워크시트가 없습니다.")
        acc: List[str] = []
        for ws in wb.worksheets:
            try:
                if XLSX_SKIP_HIDDEN_SHEETS and getattr(ws, "sheet_state", "visible") != "visible":
                    continue
            except Exception:
                pass
            acc.append(f"\n### [Sheet] {ws.title}")
            rows = 0
            iter_kwargs = {"values_only": True}
            if XLSX_MAX_COLS_PER_SHEET and XLSX_MAX_COLS_PER_SHEET > 0:
                iter_kwargs["max_col"] = XLSX_MAX_COLS_PER_SHEET
            for row in ws.iter_rows(**iter_kwargs):
                if rows >= XLSX_MAX_ROWS_PER_SHEET:
                    acc.append(f"...(truncated at {XLSX_MAX_ROWS_PER_SHEET} rows)")
                    break
                last = -1
                for i, v in enumerate(row):
                    sv = (str(v).strip() if v is not None else "")
                    if sv != "":
                        last = i
                if last < 0:
                    continue
                width = last + 1
                if XLSX_MAX_COLS_PER_SHEET and XLSX_MAX_COLS_PER_SHEET > 0:
                    width = min(width, XLSX_MAX_COLS_PER_SHEET)
                row_vals = []
                for v in row[:width]:
                    row_vals.append("" if v is None else str(v).strip())
                acc.append(" | ".join(row_vals))
                rows += 1
        return "\n".join(acc)

    # ---------- Chroma ----------
    def get_chroma_collection(self, collection_name: Optional[str] = None):
        name = collection_name or COLLECTION_NAME
        try:
            chroma = chromadb.CloudClient(
                tenant=os.getenv("CHROMA_TENANT"),
                database=os.getenv("CHROMA_DATABASE"),
                api_key=os.getenv("CHROMA_API_KEY"),
            )
            return chroma.get_or_create_collection(name=name)
        except Exception as e:
            print(f"Chroma Cloud 초기화 오류: {str(e)}")
            print("로컬 ChromaDB로 폴백...")
            try:
                Path(CHROMA_PATH).mkdir(parents=True, exist_ok=True)
                chroma = chromadb.PersistentClient(
                    path=CHROMA_PATH, settings=Settings(anonymized_telemetry=False, is_persistent=True)
                )
                return chroma.get_or_create_collection(name=name)
            except Exception as e2:
                print(f"로컬 ChromaDB 초기화도 실패: {str(e2)}")
                print("인메모리 ChromaDB로 최종 폴백...")
                chroma = chromadb.Client()
                return chroma.get_or_create_collection(name=name)

    # ---------- MD에서 시트·페이지 정보 추출 ----------
    def _split_md_by_sheet(self, md_text: str) -> List[Tuple[str, str]]:
        matches = list(SHEET_HEADER_RE.finditer(md_text))
        if not matches:
            return [("unknown", md_text)]
        blocks: List[Tuple[str, str]] = []
        for i, m in enumerate(matches):
            sheet_name = m.group('name').strip()
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(md_text)
            block = md_text[start:end].strip()
            blocks.append((sheet_name, block))
        return blocks

    def _infer_page_range_in_chunk(self, chunk: str) -> Optional[str]:
        pages = [int(x) for x in PAGE_MARK_RE.findall(chunk)]
        if not pages:
            return None
        return f"{min(pages)}-{max(pages)}" if len(pages) > 1 else str(pages[0])

    # ---------- 단일 파일 인덱싱 (메타 확장 + MD 정규화 경로) ----------
    def ingest_single_file_with_metadata(
        self,
        file_path: str,
        *,
        collection_name: str,
        extra_meta: Dict[str, Any],
        show_preview: bool = True
    ) -> Tuple[int, bool]:
        print(f"📂 입력 파일: {file_path} (collection={collection_name})")
        try:
            # 1) MD 정규화(공통 경로)
            if NORMALIZE_TO_MD:
                md_text, md_meta = self.to_markdown(file_path, verbose=True)
                raw_text = md_text
                if not raw_text.strip():
                    print(f"❌ 빈 파일이거나 MD 변환 실패: {file_path}")
                    return 0, False
            else:
                # (백업 경로) 원문 텍스트
                raw_text = self._fallback_load_text(file_path)
                md_meta = {"normalizer_ver": "none", "source_type": "unknown"}
                if not raw_text.strip():
                    print(f"❌ 빈 파일이거나 읽기 실패: {file_path}")
                    return 0, False

            print(f"✅ 파일 로드 완료, 전체 길이: {len(raw_text):,} chars")

            # 2) 청킹 (XLSX-MD는 시트 우선 분할)
            chunks: List[str] = []
            chunk_sheets: List[Optional[str]] = []

            if NORMALIZE_TO_MD and md_meta.get("source_type") == "xlsx":
                for sheet_name, block in self._split_md_by_sheet(raw_text):
                    parts = self.text_splitter.split_text(block)
                    chunks.extend(parts)
                    chunk_sheets.extend([sheet_name] * len(parts))
            else:
                parts = self.text_splitter.split_text(raw_text)
                chunks = parts
                chunk_sheets = [None] * len(parts)

            if not chunks:
                print("❌ 청킹 결과가 비어 있습니다.")
                return 0, False

            if show_preview:
                for i, c in enumerate(chunks[:3]):
                    print(f"  [Chunk {i}] {len(c):,} chars / preview: {c[:120]}...")

            # 3) 임베딩 생성
            print("⚙️ 임베딩 생성 중...")
            embeddings = embed_texts_batched(chunks)
            if not embeddings:
                print("❌ 임베딩 생성 실패(빈 입력).")
                return 0, False
            print(f"✅ 임베딩 완료 → shape: {len(embeddings)} x {len(embeddings[0])}")

            # 4) 저장(회사별 컬렉션)
            collection = self.get_chroma_collection(collection_name)
            file_name = Path(file_path).name

            # (증분) 같은 문서(doc_id 있으면 doc_id 단위, 없으면 source) 삭제
            try:
                if extra_meta and "doc_id" in extra_meta:
                    existing = collection.get(where={"doc_id": extra_meta["doc_id"]})
                else:
                    existing = collection.get(where={"source": file_name})
                if existing and existing.get("ids"):
                    collection.delete(ids=existing["ids"])
                    print(f"🗑 기존 {len(existing['ids'])} 청크 삭제")
            except Exception as e:
                print(f"⚠️ 기존 청크 삭제 중 오류: {e}")

            base_id = Path(file_path).stem
            ids = [f"{base_id}-{i}" for i in range(len(chunks))]

            # 5) 메타 확장
            file_hash = _sha256(raw_text)  # 간단 파일 해시(정규화 결과 기준)
            metadatas = []
            for i, c in enumerate(chunks):
                chunk_hash = _sha256(c)
                chunk_tokens = _estimate_tokens(c)
                overlap_used = CHUNK_OVERLAP  # 현 단계는 고정
                page_range = self._infer_page_range_in_chunk(c) if md_meta.get("source_type") == "pdf" else None
                m = {
                    # 기본
                    "source": file_name,
                    "chunk_idx": i,
                    "content_type": md_meta.get("source_type", "unknown"),
                    "normalizer_ver": md_meta.get("normalizer_ver", "none"),
                    # 메타 확장
                    "file_hash": file_hash,
                    "chunk_hash": chunk_hash,
                    "chunk_tokens": chunk_tokens,
                    "overlap_used": overlap_used,
                    "sheet": chunk_sheets[i] if chunk_sheets[i] is not None else "",
                    "page": page_range if page_range is not None else "",
                    # 선택 전달 메타
                }
                if isinstance(extra_meta, dict):
                    m.update(extra_meta)
                metadatas.append(m)

            collection.add(ids=ids, metadatas=metadatas, embeddings=embeddings, documents=chunks)
            print(f"🎉 완료! {len(chunks)} chunks → Chroma collection '{collection_name}' 저장")
            return len(chunks), True

        except Exception as e:
            print(f"❌ 파일 처리 중 오류 발생: {str(e)}")
            return 0, False

    # (백업) 원문 로더 — 정상 경로는 to_markdown()
    def _fallback_load_text(self, file_path: str) -> str:
        p = Path(file_path); ext = p.suffix.lower(); actual = _detect_office_kind(p)
        try:
            if ext == ".pdf":
                return self.read_pdf(p)
            if ext == ".docx" or (actual == "docx" and ext != ".xlsx"):
                return self.read_docx(p)
            if ext == ".xlsx" or (actual == "xlsx" and ext != ".docx"):
                return self.read_xlsx(p)
            if ext == ".csv":
                return self.read_txt(p)  # 간단: CSV -> 텍스트 (필요시 별도 구현)
            if ext == ".txt":
                return self.read_txt(p)
            if actual == "docx":
                return self.read_docx(p)
            if actual == "xlsx":
                return self.read_xlsx(p)
            return ""
        except Exception:
            return ""

    # 간단 텍스트 리더(인코딩 폴백)
    def read_txt(self, path: Path) -> str:
        encodings = ['utf-8-sig', 'utf-8', 'cp949', 'euc-kr', 'latin1']
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                    if content.strip():
                        return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError(f"텍스트 파일 인코딩을 감지할 수 없습니다: {path}")

# -------------------- CLI --------------------
def main():
    print("=" * 80)
    print("📚 문서 임베딩 서비스 (MD 정규화 + 시트 단위 청킹 + 메타 확장)")
    print("=" * 80)

    if len(sys.argv) < 2:
        print("사용법:")
        print("  테스트용 파일 처리: python internal_ingest.py <파일경로> [<컬렉션명>]")
        print("예시:")
        print("  python internal_ingest.py ./storage/data/document.pdf inside_data")
        sys.exit(1)

    path = sys.argv[1]
    collection = sys.argv[2] if len(sys.argv) > 2 else COLLECTION_NAME

    try:
        path_obj = Path(path)
        if path_obj.is_file():
            svc = IngestService()
            cnt, ok = svc.ingest_single_file_with_metadata(
                str(path_obj),
                collection_name=collection,
                extra_meta={},  # 필요 시 doc_id/company_id/user_id/is_private 등 주입
                show_preview=True
            )
            sys.exit(0 if ok else 1)
        else:
            print(f"❌ 경로가 존재하지 않습니다: {path}")
            sys.exit(1)
    except KeyboardInterrupt:
        print("\n\n❌ 사용자에 의해 중단되었습니다.")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ 예상치 못한 오류 발생: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
